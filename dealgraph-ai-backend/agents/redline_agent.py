import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from tools.audit_log import log_audit_event

def generate_redlined_contract(deal_data: dict) -> str:
    """
    Generates a mock .docx file representing a redlined counter-proposal.
    Changes unacceptable terms (e.g., Net-90) back to standard terms (Net-30) and highlights them.
    """
    filename = f"Redlined_Amended_Draft_{deal_data.get('customer_name', 'Deal').replace(' ', '_')}.docx"
    filepath = os.path.join(os.getcwd(), filename)
    
    doc = Document()
    doc.add_heading(f"Amended Draft: Master Services Agreement for {deal_data.get('customer_name')}", 0)
    
    doc.add_paragraph(f"Deal Value: ${deal_data.get('deal_value', 0):,.2f}")
    doc.add_paragraph(f"Contract Term: {deal_data.get('contract_term_months', 12)} months")
    
    doc.add_heading("Redlined Terms", level=1)
    
    # Simulate Redlining Logic
    payment_terms = deal_data.get("payment_terms", "Net-30")
    if "90" in payment_terms or "60" in payment_terms:
        p = doc.add_paragraph("Payment Terms: ")
        
        # Strikethrough the old term
        run1 = p.add_run(f"[{payment_terms}] ")
        run1.font.strike = True
        run1.font.color.rgb = RGBColor(255, 0, 0) # Red
        
        # Add the new term
        run2 = p.add_run("[Net-30]")
        run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 128, 0) # Green
    else:
        doc.add_paragraph(f"Payment Terms: {payment_terms}")

    discount = deal_data.get("discount_requested", 0.0)
    if discount > 20:
        p = doc.add_paragraph("Discount: ")
        run1 = p.add_run(f"[{discount}%] ")
        run1.font.strike = True
        run1.font.color.rgb = RGBColor(255, 0, 0)
        
        run2 = p.add_run("[20.0% MAX]")
        run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 128, 0)
    else:
        doc.add_paragraph(f"Discount: {discount}%")
        
    doc.add_paragraph("\n[AI AGENT NOTE: Generated autonomously by DealGraph Legal Agent. Ready for DocuSign routing.]")
    
    doc.save(filepath)
    
    log_audit_event(
        deal_id=deal_data.get("id", "Unknown"),
        agent_name="Legal_Agent",
        action="Generated Redlined Contract",
        result=f"Saved to {filepath}"
    )
    
    return filepath
